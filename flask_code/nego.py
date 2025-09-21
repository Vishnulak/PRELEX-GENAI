from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import logging
import json
import re
from typing import List, Tuple, Dict, Any, Optional
from dotenv import load_dotenv
import traceback
import PyPDF2
import docx
from io import BytesIO
from google.cloud import documentai
import vertexai
from vertexai.generative_models import GenerativeModel

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# CORS configuration
CORS(app, 
     origins=["http://localhost:*", "http://127.0.0.1:*", "https://localhost:*", "*"],
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
     allow_headers=["Content-Type", "Authorization", "Accept", "Origin", "X-Requested-With"],
     supports_credentials=True)

# Configuration
PROJECT_ID = os.getenv('PROJECT_ID')
LOCATION = os.getenv('LOCATION', 'us')
VERTEX_LOCATION = os.getenv('VERTEX_LOCATION', 'us-central1')
PROCESSOR_ID = os.getenv('PROCESSOR_ID')
MODEL_NAME = os.getenv('MODEL_NAME', 'gemini-2.5-flash-lite')

# File processing limits
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
MAX_DOC_CHARS = 50000
MAX_TEXT_LENGTH = 100000
MIN_TEXT_LENGTH = 50

# Supported file types
SUPPORTED_MIME_TYPES = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc'
}

EXTENSION_TO_MIME = {
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.doc': 'application/msword'
}

# Global clients
_document_ai_client: Optional[documentai.DocumentProcessorServiceClient] = None
_vertex_ai_initialized = False
_vertex_ai_available = False

# Try to import python-magic with graceful fallback
try:
    import magic
    MAGIC_AVAILABLE = True
    logger.info("python-magic library loaded successfully")
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning("python-magic not available. Using fallback MIME detection.")

# COMPREHENSIVE RISK PATTERNS
COMPREHENSIVE_RISK_PATTERNS = {
    'termination_without_cause': {
        'patterns': [
            r'(?i)terminat(?:e|ion).*without.*(?:cause|reason|notice)',
            r'(?i)(?:may|can|shall).*terminat(?:e|ion).*(?:at.*will|immediately|discretion)',
            r'(?i)right.*to.*terminat(?:e|ion).*for.*any.*reason',
            r'(?i)end.*(?:this.*agreement|contract).*without.*(?:cause|reason)',
            r'(?i)cancel(?:lation)?.*without.*(?:notice|cause|penalty)',
        ],
        'severity': 'high',
        'category': 'termination',
        'plain_english': 'They can cancel your contract anytime without giving you a reason',
        'hidden_tricks': [
            'No advance warning required - you could lose everything instantly',
            'They keep all payments made, even for unused services',
            'You have no recourse or compensation for sudden termination',
            'Often paired with non-refund clauses to maximize their protection'
        ],
        'real_world_consequences': [
            'Immediate loss of services you\'ve paid for',
            'Business disruption if you depend on their services',
            'Lost time and money finding replacement services',
            'Potential legal costs with no guarantee of recovery'
        ],
        'negotiation_tips': [
            'Demand 30-60 days written notice minimum',
            'Require them to state a valid business reason',
            'Negotiate pro-rated refund for unused services',
            'Add penalty clause if they terminate without cause',
            'Include transition assistance provision'
        ],
        'comparative_justice': 'Fair contracts require 30+ days notice and valid cause. Consumer protection laws in many states require reasonable notice.',
        'red_flags': ['immediate termination', 'sole discretion', 'any reason', 'without notice']
    },
    
    'unlimited_liability': {
        'patterns': [
            r'(?i)unlimited.*liability',
            r'(?i)liable.*for.*all.*(?:damages|costs|losses)',
            r'(?i)indemnify.*(?:and.*hold.*harmless|defend).*(?:from|against).*(?:all|any).*claims',
            r'(?i)personal.*guarantee.*for.*all',
            r'(?i)jointly.*and.*severally.*liable',
            r'(?i)responsible.*for.*all.*(?:legal.*costs|attorney.*fees|damages).*arising'
        ],
        'severity': 'critical',
        'category': 'financial_liability',
        'plain_english': 'You are responsible for unlimited damages and costs if anything goes wrong',
        'hidden_tricks': [
            'No cap on how much you could owe - could be millions',
            'You pay even if the problem wasn\'t your fault',
            'Includes their legal fees, not just damages',
            'May apply to actions of your employees or contractors',
            'Could affect your personal assets, not just business'
        ],
        'real_world_consequences': [
            'Bankruptcy risk from unlimited financial exposure',
            'Personal assets at risk (house, savings, retirement)',
            'Credit destruction from judgments',
            'Inability to get future contracts or loans',
            'Family financial security threatened'
        ],
        'negotiation_tips': [
            'Cap total liability at contract value or reasonable amount',
            'Exclude liability for their negligence or willful misconduct',
            'Require mutual indemnification (both parties protect each other)',
            'Add insurance requirements instead of unlimited liability',
            'Limit liability to direct damages only, exclude consequential'
        ],
        'comparative_justice': 'Standard business practice limits liability to contract amount. Unlimited liability is predatory and often unenforceable.',
        'red_flags': ['unlimited', 'all damages', 'joint and several', 'personal guarantee']
    },
    
    'automatic_renewal_trap': {
        'patterns': [
            r'(?i)automatic(?:ally)?.*renew(?:al|s)?',
            r'(?i)renew(?:s|al).*automatic(?:ally)?.*unless.*(?:cancelled|terminated)',
            r'(?i)contract.*continues.*unless.*written.*notice',
            r'(?i)evergreen.*clause',
            r'(?i)perpetual.*renewal',
            r'(?i)notice.*(?:30|60|90).*days.*prior.*to.*renewal'
        ],
        'severity': 'medium-high',
        'category': 'contract_terms',
        'plain_english': 'Your contract automatically extends and charges you again unless you actively cancel',
        'hidden_tricks': [
            'Short cancellation windows (often 30-90 days before renewal)',
            'Cancellation must be in writing, not just verbal',
            'New terms can be imposed with each renewal',
            'Price increases often take effect with renewal',
            'Forgetting to cancel locks you in for another full term'
        ],
        'real_world_consequences': [
            'Unexpected charges on your credit card or bank account',
            'Locked into services you no longer need',
            'Difficulty canceling once auto-renewed',
            'Compounding costs over multiple renewal cycles',
            'Legal obligation to pay even if service quality declines'
        ],
        'negotiation_tips': [
            'Change to manual renewal requiring your active consent',
            'Extend cancellation notice period to 90+ days',
            'Require email reminders before renewal deadlines',
            'Allow cancellation at any time with pro-rated refund',
            'Lock in current pricing for future renewals'
        ],
        'comparative_justice': 'Consumer-friendly contracts require opt-in renewal. Auto-renewal should have generous cancellation periods and clear notifications.',
        'red_flags': ['automatic renewal', 'evergreen', 'unless cancelled', 'perpetual']
    },
    
    'binding_arbitration': {
        'patterns': [
            r'(?i)binding.*arbitration',
            r'(?i)disputes.*(?:must|shall).*be.*(?:resolved|settled).*(?:through|by).*arbitration',
            r'(?i)waive.*right.*to.*(?:jury|court|trial)',
            r'(?i)exclusive.*jurisdiction.*arbitration',
            r'(?i)class.*action.*waiver',
            r'(?i)mandatory.*arbitration'
        ],
        'severity': 'high',
        'category': 'legal_rights',
        'plain_english': 'You give up your right to sue them in court and must use private arbitration',
        'hidden_tricks': [
            'Arbitrator is often chosen/paid by the company',
            'No jury of your peers - single arbitrator decides',
            'Limited ability to appeal unfavorable decisions',
            'Discovery process is restricted (less evidence allowed)',
            'Often combined with class action waivers'
        ],
        'real_world_consequences': [
            'Loss of constitutional right to jury trial',
            'Higher costs for individual arbitration vs. court',
            'Arbitrators may favor repeat corporate clients',
            'Limited public record of disputes and outcomes',
            'Cannot join with other victims in class action'
        ],
        'negotiation_tips': [
            'Require mediation before arbitration',
            'Allow court option for claims under $10,000',
            'Mutually select neutral arbitrator',
            'Share arbitration costs equally',
            'Preserve right to seek injunctive relief in court'
        ],
        'comparative_justice': 'Many states restrict forced arbitration. Supreme Court has limited some arbitration requirements in consumer contracts.',
        'red_flags': ['binding arbitration', 'waive right to court', 'class action waiver']
    },
    
    'liquidated_damages_penalty': {
        'patterns': [
            r'(?i)liquidated.*damages.*(?:of|equal.*to|\$)',
            r'(?i)penalty.*(?:of|equal.*to).*\$[\d,]+',
            r'(?i)forfeit.*(?:deposit|payment|fee).*(?:of|totaling)',
            r'(?i)damages.*(?:equal.*to|of).*(?:\d+.*times|multiple.*of)',
            r'(?i)punitive.*damages.*(?:of|\$)',
            r'(?i)breach.*results.*in.*payment.*of.*\$[\d,]+'
        ],
        'severity': 'medium-high',
        'category': 'financial_penalties',
        'plain_english': 'You must pay specific penalty amounts for breaking any part of the contract',
        'hidden_tricks': [
            'Penalties often far exceed actual damages',
            'Apply to minor technical breaches, not just major ones',
            'No consideration of your ability to pay',
            'May be triggered by circumstances beyond your control',
            'Often non-negotiable once contract is signed'
        ],
        'real_world_consequences': [
            'Large financial penalties for minor violations',
            'Double punishment (lose service AND pay penalty)',
            'Debt collection and credit damage if unpaid',
            'Legal costs to dispute unreasonable penalties',
            'Business cash flow problems from unexpected penalties'
        ],
        'negotiation_tips': [
            'Ensure penalties reflect reasonable estimate of actual damages',
            'Add materiality threshold (only for significant breaches)',
            'Require notice and cure period before penalties apply',
            'Cap penalties at reasonable percentage of contract value',
            'Make penalties mutual (they pay you if they breach too)'
        ],
        'comparative_justice': 'Courts may refuse to enforce penalties that are grossly disproportionate to actual damages.',
        'red_flags': ['liquidated damages', 'penalty of $', 'forfeit', 'punitive damages']
    },
    
    'non_refundable_trap': {
        'patterns': [
            r'(?i)non-?refundable',
            r'(?i)no.*refund(?:s)?.*(?:under|in).*any.*circumstance',
            r'(?i)all.*(?:payments|fees).*are.*final',
            r'(?i)deposits?.*(?:are|will.*be).*(?:retained|kept|forfeited)',
            r'(?i)payment.*not.*returnable',
            r'(?i)fees.*paid.*in.*advance.*non-?refundable'
        ],
        'severity': 'medium',
        'category': 'payment_terms',
        'plain_english': 'You cannot get your money back under any circumstances, even if they fail to deliver',
        'hidden_tricks': [
            'No refunds even if they breach the contract',
            'No refunds for services never provided',
            'No refunds if they go out of business',
            'May apply to large advance payments or deposits',
            'Often buried in fine print or addendums'
        ],
        'real_world_consequences': [
            'Total loss of advance payments if service fails',
            'No recourse if company fails to perform',
            'Incentivizes company to take payment without delivering',
            'Financial loss even if you have legitimate complaints',
            'Difficulty getting credit card chargebacks'
        ],
        'negotiation_tips': [
            'Negotiate partial refunds for undelivered services',
            'Add performance milestones tied to payment schedule',
            'Include refund provisions for their material breach',
            'Limit non-refundable amounts to actual costs incurred',
            'Add escrow for large advance payments'
        ],
        'comparative_justice': 'Consumer protection laws often override blanket non-refund clauses. Fair contracts provide refunds for non-performance.',
        'red_flags': ['non-refundable', 'no refunds', 'payments are final', 'deposits retained']
    }
}

# Additional specialized patterns for different contract types
CONTRACT_TYPE_PATTERNS = {
    'employment': {
        'non_compete_overreach': {
            'patterns': [
                r'(?i)non-?compete.*(?:for|period.*of).*(?:\d+.*years?|indefinitely)',
                r'(?i)shall.*not.*(?:compete|engage.*in.*similar.*business)',
                r'(?i)restraint.*of.*trade.*(?:for|during).*(?:\d+.*years?)',
                r'(?i)covenant.*not.*to.*compete.*(?:worldwide|nationally)'
            ],
            'severity': 'high',
            'category': 'employment_restrictions',
            'plain_english': 'You cannot work in your field for an unreasonably long time or broad area',
            'consequence': 'Loss of livelihood and career advancement opportunities'
        }
    },
    'software': {
        'data_harvesting': {
            'patterns': [
                r'(?i)collect.*all.*(?:data|information|analytics).*generated',
                r'(?i)right.*to.*use.*(?:customer.*data|user.*information).*for.*any.*purpose',
                r'(?i)license.*to.*use.*(?:your.*data|information.*provided)',
                r'(?i)aggregate.*(?:data|information).*for.*(?:commercial|business).*purposes'
            ],
            'severity': 'medium-high',
            'category': 'data_privacy',
            'plain_english': 'They can collect and sell your business data and customer information',
            'consequence': 'Loss of competitive advantage and customer privacy'
        }
    }
}

def initialize_services():
    """Initialize Document AI and Vertex AI services"""
    global _document_ai_client, _vertex_ai_initialized, _vertex_ai_available
    
    # Initialize Document AI
    if PROJECT_ID and PROCESSOR_ID:
        try:
            _document_ai_client = documentai.DocumentProcessorServiceClient()
            logger.info("Document AI client initialized successfully")
        except Exception as e:
            logger.warning(f"Failed to initialize Document AI: {e}")
            _document_ai_client = None
    
    # Initialize Vertex AI
    if PROJECT_ID and not _vertex_ai_initialized:
        try:
            vertexai.init(project=PROJECT_ID, location=VERTEX_LOCATION)
            _vertex_ai_initialized = True
            _vertex_ai_available = True
            logger.info(f"Vertex AI initialized successfully")
        except Exception as e:
            _vertex_ai_initialized = True
            _vertex_ai_available = False
            logger.warning(f"Failed to initialize Vertex AI: {e}")

def detect_mime_type(file_content: bytes, filename: str) -> str:
    """Detect MIME type using multiple methods"""
    # Method 1: Try python-magic if available
    if MAGIC_AVAILABLE:
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
            if mime_type in SUPPORTED_MIME_TYPES:
                return mime_type
        except Exception as e:
            logger.warning(f"python-magic detection failed: {e}")
    
    # Method 2: Check file signature (magic bytes)
    if file_content.startswith(b'%PDF'):
        return 'application/pdf'
    elif file_content.startswith(b'PK\x03\x04') or file_content.startswith(b'PK\x05\x06') or file_content.startswith(b'PK\x07\x08'):
        if filename and filename.lower().endswith('.docx'):
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif file_content.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
        if filename and filename.lower().endswith('.doc'):
            return 'application/msword'
    
    # Method 3: Extension-based fallback
    if filename:
        _, ext = os.path.splitext(filename.lower())
        if ext in EXTENSION_TO_MIME:
            return EXTENSION_TO_MIME[ext]
    
    raise ValueError(f"Could not determine MIME type for file: {filename}")

def extract_text_fallback(file_content: bytes, mime_type: str) -> str:
    """Fallback text extraction using local libraries"""
    text = ""
    
    if mime_type == 'application/pdf':
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                text += f"\n--- Page {page_num + 1} (extraction failed) ---\n"
            
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = docx.Document(BytesIO(file_content))
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    
    if not text.strip():
        raise Exception("No text could be extracted from the document")
    
    return text.strip()

def extract_text_with_document_ai(file_content: bytes, mime_type: str) -> str:
    """Extract text using Document AI or fallback"""
    try:
        if not _document_ai_client:
            logger.info("Document AI not available, using fallback extraction")
            return extract_text_fallback(file_content, mime_type)
        
        name = _document_ai_client.processor_path(PROJECT_ID, LOCATION, PROCESSOR_ID)
        
        request_obj = documentai.ProcessRequest(
            name=name,
            raw_document=documentai.RawDocument(
                content=file_content,
                mime_type=mime_type
            )
        )
        
        result = _document_ai_client.process_document(request=request_obj)
        extracted_text = result.document.text if result.document.text else ""
        
        if not extracted_text.strip():
            logger.warning("Document AI returned empty text, using fallback")
            return extract_text_fallback(file_content, mime_type)
        
        if len(extracted_text) > MAX_DOC_CHARS:
            extracted_text = extracted_text[:MAX_DOC_CHARS]
        
        logger.info(f"Document AI extraction successful: {len(extracted_text)} characters")
        return extracted_text
        
    except Exception as e:
        logger.error(f"Document AI extraction failed: {e}, using fallback")
        return extract_text_fallback(file_content, mime_type)

def detect_contract_type(text: str) -> str:
    """Detect the type of contract to apply specialized patterns"""
    text_lower = text.lower()
    
    if any(keyword in text_lower for keyword in ['employee', 'employment', 'job', 'position', 'salary', 'wages']):
        return 'employment'
    elif any(keyword in text_lower for keyword in ['software', 'application', 'development', 'coding', 'programming']):
        return 'software'
    elif any(keyword in text_lower for keyword in ['lease', 'rent', 'tenant', 'landlord', 'property']):
        return 'rental'
    elif any(keyword in text_lower for keyword in ['service', 'consulting', 'professional services']):
        return 'service'
    elif any(keyword in text_lower for keyword in ['purchase', 'sale', 'buy', 'sell', 'goods']):
        return 'sales'
    else:
        return 'general'

def extract_key_information_enhanced(text: str) -> Dict[str, Any]:
    """Enhanced key information extraction with deeper analysis"""
    key_info = {
        "parties": [],
        "dates": [],
        "amounts": [],
        "obligations": [],
        "rights": [],
        "termination_clauses": [],
        "payment_terms": [],
        "penalty_clauses": [],
        "liability_clauses": [],
        "confidentiality_clauses": [],
        "modification_clauses": [],
        "governing_law": [],
        "dispute_resolution": []
    }
    
    # Enhanced party extraction
    party_patterns = [
        r'between\s+([^,\(]+(?:\([^)]+\))?)\s+(?:and|&)',
        r'(?:Client|Customer|Buyer|Tenant|Lessee|Contractor|Employee)[:\s]+([^,\.\n]+)',
        r'(?:Company|Provider|Seller|Landlord|Lessor|Employer)[:\s]+([^,\.\n]+)',
        r'(?:Corp\.|Corporation|LLC|Ltd\.?|Inc\.?)[,\s]*([^,\.\n]+)',
        r'"([^"]+)"[,\s]+(?:a|an)\s+(?:corporation|company|LLC)',
    ]
    
    for pattern in party_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            clean_match = re.sub(r'\s*\([^)]*\)', '', match).strip()
            if clean_match and len(clean_match) > 2 and clean_match not in key_info["parties"]:
                key_info["parties"].append(clean_match)
    
    # Enhanced date extraction
    date_patterns = [
        r'(?:dated?|effective|starting|begins?|ends?|expires?|due|term.*(?:begins|ends))\s+([A-Za-z]+ \d{1,2},? \d{4})',
        r'(?:on|by|before|after|until|from)\s+([A-Za-z]+ \d{1,2},? \d{4})',
        r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{4})\b',
        r'(?:term.*of|period.*of|duration.*of)\s+(\d+\s+(?:years?|months?|days?))',
    ]
    
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        key_info["dates"].extend(matches)
    
    # Enhanced monetary amount extraction
    amount_patterns = [
        r'\$[\d,]+(?:\.\d{2})?',
        r'(?:fee|cost|price|amount|payment|salary|wage|penalty|fine|deposit)\s+(?:of\s+)?\$?([\d,]+(?:\.\d{2})?)',
        r'(?:dollars?|USD)\s+([\d,]+(?:\.\d{2})?)',
        r'(?:total|sum|aggregate)\s+(?:of\s+)?\$?([\d,]+(?:\.\d{2})?)',
    ]
    
    for pattern in amount_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        key_info["amounts"].extend(matches)
    
    return key_info

def analyze_risks_with_enhanced_vertex_ai(text: str, summary_text: str, contract_type: str) -> List[Dict[str, Any]]:
    """Enhanced AI-powered risk analysis with contract type awareness"""
    if not _vertex_ai_available:
        return analyze_risks_with_enhanced_rules(text, summary_text, contract_type)
    
    try:
        # Create specialized prompt based on contract type
        contract_specific_guidance = ""
        if contract_type == 'employment':
            contract_specific_guidance = """
            Pay special attention to:
            - Non-compete clauses and geographic/time restrictions
            - Wage and hour provisions, overtime exemptions
            - Intellectual property assignments
            - At-will employment modifications
            - Benefits and severance terms
            """
        elif contract_type == 'software':
            contract_specific_guidance = """
            Pay special attention to:
            - Data usage and privacy rights
            - Source code ownership and licensing
            - Service level agreements and uptime guarantees
            - Limitation of liability for software defects
            - Automatic updates and feature changes
            """
        
        prompt = f"""
You are an expert legal analyst specializing in protecting consumers and small businesses from predatory contract terms. 

CONTRACT TYPE: {contract_type.title()}
{contract_specific_guidance}

DOCUMENT TEXT (first 12000 chars):
{text[:12000]}

SUMMARY:
{summary_text[:2000]}

Your task is to identify truly problematic clauses that could harm the person signing this contract. Focus on HIDDEN DANGERS and PREDATORY TERMS.

Provide a JSON response with this exact structure:

{{
  "risky_clauses": [
    {{
      "clause_number": 1,
      "clause_text": "exact problematic text from document (keep under 200 chars)",
      "plain_english": "simple explanation without legal jargon",
      "hidden_tricks": [
        "specific deceptive aspect 1",
        "specific deceptive aspect 2"
      ],
      "real_world_consequences": [
        "concrete consequence 1",
        "concrete consequence 2"
      ],
      "negotiation_tips": [
        "specific actionable tip 1",
        "specific actionable tip 2"
      ],
      "comparative_justice": "how this compares to fair/standard industry practice",
      "severity": "critical|high|medium-high|medium|low",
      "risk_category": "financial|legal_rights|termination|liability|privacy|employment|other",
      "red_flags": ["key warning phrase 1", "key warning phrase 2"]
    }}
  ]
}}

REQUIREMENTS:
- Only include genuinely risky/unfair clauses
- Be very specific about hidden tricks and consequences
- Provide actionable negotiation advice
- Compare to industry standards
- Use simple language anyone can understand
- Maximum 10 clauses, prioritize worst ones
- Include exact text quotes from the document
"""

        model = GenerativeModel(MODEL_NAME)
        generation_config = {
            "temperature": 0.1,
            "max_output_tokens": 4000,
            "top_p": 0.8,
            "top_k": 20
        }
        
        response = model.generate_content(prompt, generation_config=generation_config)
        response_text = response.text.strip()
        
        # Extract JSON
        if '```json' in response_text:
            json_start = response_text.find('```json') + 7
            json_end = response_text.find('```', json_start)
            json_text = response_text[json_start:json_end].strip()
        elif '{' in response_text and '}' in response_text:
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            json_text = response_text[json_start:json_end]
        else:
            raise ValueError("No JSON found in response")
        
        result = json.loads(json_text)
        return result.get('risky_clauses', [])
        
    except Exception as e:
        logger.error(f"Enhanced Vertex AI risk analysis failed: {e}")
        return analyze_risks_with_enhanced_rules(text, summary_text, contract_type)

def analyze_risks_with_enhanced_rules(text: str, summary_text: str, contract_type: str) -> List[Dict[str, Any]]:
    """Enhanced rule-based risk analysis with contract-type specific patterns"""
    risky_clauses = []
    full_text = f"{text}\n\n{summary_text}"
    clause_number = 1
    
    # Combine general patterns with contract-specific patterns
    all_patterns = dict(COMPREHENSIVE_RISK_PATTERNS)
    if contract_type in CONTRACT_TYPE_PATTERNS:
        all_patterns.update(CONTRACT_TYPE_PATTERNS[contract_type])
    
    # Sort patterns by severity
    severity_order = {'critical': 0, 'high': 1, 'medium-high': 2, 'medium': 3, 'low': 4}
    sorted_patterns = sorted(all_patterns.items(), 
                           key=lambda x: severity_order.get(x[1].get('severity', 'medium'), 3))
    
    for risk_type, risk_data in sorted_patterns:
        patterns = risk_data['patterns']
        found_match = False
        
        for pattern in patterns:
            matches = list(re.finditer(pattern, full_text, re.IGNORECASE | re.MULTILINE))
            
            for match in matches:
                if found_match:
                    continue
                
                # Extract context around match
                start = max(0, match.start() - 100)
                end = min(len(full_text), match.end() + 100)
                
                context = full_text[start:end]
                sentences = re.split(r'[.!?]+', context)
                
                clause_text = ' '.join(sentences[max(0, len(sentences)//2-1):len(sentences)//2+2]).strip()
                clause_text = ' '.join(clause_text.split())
                if len(clause_text) > 200:
                    clause_text = clause_text[:197] + "..."
                
                risky_clauses.append({
                    "clause_number": clause_number,
                    "clause_text": clause_text,
                    "plain_english": risk_data['plain_english'],
                    "hidden_tricks": risk_data.get('hidden_tricks', [risk_data['plain_english']]),
                    "real_world_consequences": risk_data.get('real_world_consequences', [risk_data.get('consequence', 'Could result in financial or legal problems')]),
                    "negotiation_tips": risk_data.get('negotiation_tips', [risk_data.get('negotiation_tip', 'Negotiate better terms or seek legal advice')]),
                    "comparative_justice": risk_data['comparative_justice'],
                    "severity": risk_data.get('severity', 'medium'),
                    "risk_category": risk_data.get('category', 'general'),
                    "red_flags": risk_data.get('red_flags', [])
                })
                
                clause_number += 1
                found_match = True
                
                if clause_number > 10:
                    break
                    
        if clause_number > 10:
            break
    
    return risky_clauses

def generate_enhanced_summary_with_vertex_ai(text: str, key_info: Dict[str, Any], contract_type: str) -> str:
    """Generate enhanced summary with contract type awareness"""
    if not _vertex_ai_available:
        return generate_enhanced_fallback_summary(text, key_info, contract_type)
    
    try:
        prompt = f"""
You are a professional legal analyst helping ordinary people understand complex legal documents.

CONTRACT TYPE: {contract_type.title()}

DOCUMENT TEXT:
{text[:MAX_TEXT_LENGTH]}

KEY INFORMATION:
- Parties: {', '.join(key_info.get('parties', [])[:3])}
- Dates: {', '.join(key_info.get('dates', [])[:3])}
- Amounts: {', '.join(key_info.get('amounts', [])[:5])}

Create a comprehensive summary with these sections:

**Document Overview**
Explain what type of agreement this is and its purpose.

**Parties and Relationships** 
Describe the main parties and their roles.

**Key Terms and Obligations**
Cover main services, financial commitments, timelines, and performance expectations.

**Rights and Responsibilities**
Explain what each party gets and must provide.

**Important Provisions**
Describe termination, penalties, dispute resolution, and modification terms.

**Professional Assessment**
Provide balanced evaluation including potential risks and recommendations.

Write in clear, professional language without legal jargon. Use flowing paragraphs, not bullet points.
"""

        model = GenerativeModel(MODEL_NAME)
        generation_config = {
            "temperature": 0.3,
            "max_output_tokens": 2500,
            "top_p": 0.9,
            "top_k": 40
        }
        
        response = model.generate_content(prompt, generation_config=generation_config)
        return response.text.strip()
        
    except Exception as e:
        logger.error(f"Vertex AI summary failed: {e}")
        return generate_enhanced_fallback_summary(text, key_info, contract_type)

def generate_enhanced_fallback_summary(text: str, key_info: Dict[str, Any], contract_type: str) -> str:
    """Generate fallback summary with contract type awareness"""
    contract_descriptions = {
        'employment': 'job or employment contract that defines your work relationship',
        'software': 'software service agreement that governs your use of technology services',
        'rental': 'rental or lease agreement for property or equipment',
        'service': 'professional service contract for specific work to be performed',
        'sales': 'purchase or sales agreement for goods or products',
        'general': 'legal agreement that creates binding obligations'
    }
    
    summary_parts = []
    
    summary_parts.append("**Document Overview**")
    summary_parts.append(f"This is a {contract_descriptions.get(contract_type, 'legal agreement')} between the parties listed below. ")
    summary_parts.append(f"The agreement establishes terms for {contract_type.replace('_', ' ')} services and defines rights and obligations.")
    summary_parts.append("")
    
    summary_parts.append("**Parties and Relationships**")
    if key_info["parties"]:
        parties_text = "The main parties are: " + ", ".join(key_info["parties"][:4]) + ". "
        summary_parts.append(parties_text + "Each party has specific roles and responsibilities.")
    else:
        summary_parts.append("Multiple parties are involved with specific roles defined in the document.")
    summary_parts.append("")
    
    summary_parts.append("**Key Terms and Obligations**")
    if key_info["amounts"]:
        amounts_text = "Financial aspects include: " + ", ".join(str(amt) for amt in key_info["amounts"][:4])
        summary_parts.append(amounts_text)
    if key_info["dates"]:
        dates_text = "Important dates include: " + ", ".join(str(date) for date in key_info["dates"][:3])
        summary_parts.append(dates_text)
    summary_parts.append("")
    
    summary_parts.append("**Rights and Responsibilities**")
    summary_parts.append("Each party has defined rights and must fulfill certain obligations. ")
    summary_parts.append("Your rights include receiving agreed services. Your responsibilities include making payments and complying with terms.")
    summary_parts.append("")
    
    summary_parts.append("**Important Provisions**")
    summary_parts.append("The contract includes provisions regarding termination, dispute resolution, and modifications. ")
    if contract_type == 'employment':
        summary_parts.append("Pay attention to non-compete restrictions, overtime provisions, and termination procedures. ")
    elif contract_type == 'software':
        summary_parts.append("Review data usage rights, service levels, and automatic renewal terms. ")
    summary_parts.append("")
    
    summary_parts.append("**Professional Assessment**")
    summary_parts.append(f"This {contract_descriptions.get(contract_type, 'agreement')} creates binding legal obligations. ")
    summary_parts.append("Review terms carefully to ensure fairness. Consider legal advice for significant commitments.")
    
    return "\n".join(summary_parts)

def format_enhanced_final_analysis(risky_clauses: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Format the analysis into structured output"""
    if not risky_clauses:
        return {
            "risky_clauses": [],
            "consequences": [],
            "negotiation_points": [],
            "comparative_justice": [],
            "hidden_tricks": [],
            "detailed_clauses": [],
            "summary": {
                "total_risks": 0,
                "risk_level": "low",
                "recommendation": "This document appears to have standard terms with minimal risks.",
                "critical_risks": 0,
                "high_risks": 0,
                "medium_high_risks": 0,
                "categories_affected": []
            }
        }
    
    # Sort clauses by severity
    severity_order = {'critical': 0, 'high': 1, 'medium-high': 2, 'medium': 3, 'low': 4}
    sorted_clauses = sorted(risky_clauses, 
                           key=lambda x: (severity_order.get(x.get('severity', 'medium'), 3), 
                                        x.get('clause_number', 999)))
    
    risky_clauses_list = []
    consequences_list = []
    negotiation_points_list = []
    comparative_justice_list = []
    hidden_tricks_list = []
    
    # Count risks by severity
    critical_count = sum(1 for clause in sorted_clauses if clause.get('severity') == 'critical')
    high_count = sum(1 for clause in sorted_clauses if clause.get('severity') == 'high')
    medium_high_count = sum(1 for clause in sorted_clauses if clause.get('severity') == 'medium-high')
    
    for i, clause in enumerate(sorted_clauses, 1):
        severity_icon = {
            'critical': '🚨',
            'high': '⚠️',
            'medium-high': '⚡',
            'medium': '📋',
            'low': 'ℹ️'
        }.get(clause.get('severity', 'medium'), '📋')
        
        risky_clauses_list.append(f"{severity_icon} {i}. {clause['plain_english']}")
        
        # Handle consequences
        consequences = clause.get('real_world_consequences', clause.get('consequence', ['Potential negative impact']))
        if isinstance(consequences, str):
            consequences = [consequences]
        consequences_list.append(f"{i}. " + " | ".join(consequences[:2]))
        
        # Handle negotiation tips
        negotiation = clause.get('negotiation_tips', clause.get('negotiation_tip', ['Seek legal advice']))
        if isinstance(negotiation, str):
            negotiation = [negotiation]
        negotiation_points_list.append(f"{i}. " + " | ".join(negotiation[:2]))
        
        comparative_justice_list.append(f"{i}. {clause.get('comparative_justice', 'Review against industry standards')}")
        
        # Handle hidden tricks
        tricks = clause.get('hidden_tricks', [clause.get('plain_english', 'Hidden unfavorable term')])
        if isinstance(tricks, str):
            tricks = [tricks]
        hidden_tricks_list.append(f"{i}. " + " | ".join(tricks[:2]))
    
    # Risk assessment
    risk_count = len(sorted_clauses)
    if critical_count >= 1:
        risk_level = "critical"
        recommendation = "🚨 DANGER: Contains critical risks. DO NOT SIGN without legal review."
    elif high_count >= 3 or (high_count >= 1 and medium_high_count >= 2):
        risk_level = "high"
        recommendation = "⚠️ HIGH RISK: Multiple concerning clauses. Legal review recommended."
    elif high_count >= 1 or medium_high_count >= 3:
        risk_level = "medium-high"
        recommendation = "⚡ CAUTION: Several risky terms worth addressing."
    elif risk_count >= 2:
        risk_level = "medium"
        recommendation = "📋 REVIEW NEEDED: Some problematic clauses to negotiate."
    else:
        risk_level = "low"
        recommendation = "✅ RELATIVELY SAFE: Minimal risks identified."
    
    return {
        "risky_clauses": risky_clauses_list,
        "consequences": consequences_list,
        "negotiation_points": negotiation_points_list,
        "comparative_justice": comparative_justice_list,
        "hidden_tricks": hidden_tricks_list,
        "detailed_clauses": sorted_clauses,
        "summary": {
            "total_risks": risk_count,
            "risk_level": risk_level,
            "recommendation": recommendation,
            "critical_risks": critical_count,
            "high_risks": high_count,
            "medium_high_risks": medium_high_count,
            "categories_affected": list(set([clause.get('risk_category', 'Unknown') for clause in sorted_clauses])),
            "most_severe": sorted_clauses[0].get('severity', 'medium') if sorted_clauses else 'none'
        }
    }

# ROUTES
@app.before_request
def handle_preflight():
    if request.method == "OPTIONS":
        res = jsonify(success=True)
        res.headers['Access-Control-Allow-Origin'] = '*'
        res.headers['Access-Control-Allow-Methods'] = 'GET,POST,PUT,DELETE,OPTIONS'
        res.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization,Accept,Origin,X-Requested-With'
        return res

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization,Accept,Origin,X-Requested-With')
    response.headers.add('Access-Control-Allow-Methods', 'GET,POST,PUT,DELETE,OPTIONS')
    return response

@app.route('/analyze-document', methods=['POST', 'OPTIONS'])
def analyze_document():
    """Document analysis workflow"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        logger.info("Starting document analysis")
        
        # Validate file upload
        if 'document' not in request.files:
            return jsonify({
                'error': 'No document file provided',
                'message': 'Please upload a document'
            }), 400
        
        file = request.files['document']
        
        if not file or file.filename == '':
            return jsonify({
                'error': 'No file selected',
                'message': 'Please select a file'
            }), 400
        
        file_content = file.read()
        
        if not file_content:
            return jsonify({
                'error': 'Empty file',
                'message': 'The file appears to be empty'
            }), 400
        
        file_size = len(file_content)
        if file_size > MAX_FILE_SIZE:
            return jsonify({
                'error': 'File too large',
                'message': f'File size exceeds {MAX_FILE_SIZE:,} bytes'
            }), 413
        
        logger.info(f"Processing file: {file.filename}, size: {file_size:,} bytes")
        
        # Extract text
        try:
            mime_type = detect_mime_type(file_content, file.filename)
            logger.info(f"Detected MIME type: {mime_type}")
        except ValueError as e:
            return jsonify({
                'error': 'Unsupported file type',
                'message': str(e),
                'supported_types': list(SUPPORTED_MIME_TYPES.values())
            }), 400
        
        extracted_text = extract_text_with_document_ai(file_content, mime_type)
        
        if len(extracted_text.strip()) < MIN_TEXT_LENGTH:
            return jsonify({
                'error': 'Extracted text too short',
                'message': f'Text must be at least {MIN_TEXT_LENGTH} characters',
                'extracted_length': len(extracted_text)
            }), 400
        
        logger.info(f"Text extraction completed: {len(extracted_text)} characters")
        
        # Detect contract type
        contract_type = detect_contract_type(extracted_text)
        logger.info(f"Contract type: {contract_type}")
        
        # Extract key information
        key_info = extract_key_information_enhanced(extracted_text)
        
        # Generate summary
        summary_text = generate_enhanced_summary_with_vertex_ai(extracted_text, key_info, contract_type)
        logger.info(f"Summary generated: {len(summary_text)} characters")
        
        # Risk analysis
        risky_clauses = analyze_risks_with_enhanced_vertex_ai(extracted_text, summary_text, contract_type)
        logger.info(f"Risk analysis completed: {len(risky_clauses)} risks found")
        
        # Format analysis
        final_analysis = format_enhanced_final_analysis(risky_clauses)
        
        # Complete response
        complete_response = {
            "status": "success",
            "document_info": {
                "filename": file.filename,
                "file_size": file_size,
                "mime_type": mime_type,
                "contract_type": contract_type,
                "extracted_text_length": len(extracted_text),
                "summary_length": len(summary_text)
            },
            "extraction": {
                "text": extracted_text,
                "key_information": {
                    "parties_involved": key_info["parties"][:4],
                    "important_dates": key_info["dates"][:4],
                    "monetary_amounts": key_info["amounts"][:5],
                    "payment_terms": key_info.get("payment_terms", [])[:3],
                    "penalty_clauses": key_info.get("penalty_clauses", [])[:3],
                    "termination_info": key_info.get("termination_clauses", [])[:3]
                }
            },
            "summary": {
                "contract_type": contract_type,
                "summary_text": summary_text
            },
            "risk_analysis": {
                "risky_clauses": final_analysis["risky_clauses"],
                "hidden_tricks": final_analysis["hidden_tricks"],
                "real_world_consequences": final_analysis["consequences"],
                "negotiation_tips": final_analysis["negotiation_points"],
                "comparative_justice": final_analysis["comparative_justice"],
                "summary": final_analysis["summary"],
                "detailed_clauses": final_analysis["detailed_clauses"]
            },
            "processing_info": {
                "extraction_method": "document_ai" if _document_ai_client else "fallback",
                "summarization_method": "vertex_ai" if _vertex_ai_available else "fallback",
                "risk_analysis_method": "vertex_ai" if _vertex_ai_available else "rules_based",
                "contract_type_detected": contract_type,
                "total_risks_found": len(risky_clauses),
                "severity_breakdown": {
                    "critical": final_analysis["summary"]["critical_risks"],
                    "high": final_analysis["summary"]["high_risks"],
                    "medium_high": final_analysis["summary"]["medium_high_risks"],
                    "total": final_analysis["summary"]["total_risks"]
                }
            }
        }
        
        logger.info(f"Analysis finished: Type: {contract_type}, "
                   f"{len(extracted_text)} chars, {len(risky_clauses)} risks, "
                   f"Level: {final_analysis['summary']['risk_level']}")
        
        return jsonify(complete_response)
        
    except Exception as e:
        logger.error(f"Document analysis failed: {str(e)}")
        return jsonify({
            'error': 'Analysis failed',
            'message': str(e)
        }), 500

@app.route('/health', methods=['GET', 'OPTIONS'])
def health_check():
    """Health check endpoint"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        doc_ai_status = _document_ai_client is not None
        vertex_ai_status = _vertex_ai_available
        
        return jsonify({
            'status': 'healthy',
            'service': 'Legal Document Risk Analyzer',
            'version': '2.0.0',
            'configuration': {
                'max_file_size_mb': MAX_FILE_SIZE // (1024 * 1024),
                'supported_types': list(SUPPORTED_MIME_TYPES.values()),
                'project_id': PROJECT_ID,
                'model': MODEL_NAME
            },
            'features': {
                'document_ai_extraction': doc_ai_status,
                'vertex_ai_processing': vertex_ai_status,
                'contract_type_detection': True,
                'risk_pattern_matching': True,
                'python_magic_available': MAGIC_AVAILABLE,
                'risk_patterns': len(COMPREHENSIVE_RISK_PATTERNS)
            },
            'contract_types_supported': list(CONTRACT_TYPE_PATTERNS.keys()) + ['general'],
            'severity_levels': ['critical', 'high', 'medium-high', 'medium', 'low']
        })
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 500

@app.route('/', methods=['GET'])
def root():
    """Root endpoint"""
    return jsonify({
        'service': 'Legal Document Risk Analyzer',
        'version': '2.0.0',
        'description': 'Legal document analysis with risk detection',
        'main_endpoint': {
            'url': '/analyze-document',
            'method': 'POST',
            'description': 'Upload document for analysis',
            'input': 'Multipart form with "document" field (PDF, DOCX, DOC)'
        },
        'other_endpoints': {
            '/health': 'GET - Health check',
            '/': 'GET - API information'
        },
        'supported_file_types': list(SUPPORTED_MIME_TYPES.values()),
        'max_file_size_mb': MAX_FILE_SIZE // (1024 * 1024)
    })

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({
        'error': 'File too large',
        'message': f'Maximum file size is {MAX_FILE_SIZE:,} bytes'
    }), 413

@app.errorhandler(400)
def bad_request(error):
    return jsonify({
        'error': 'Bad request',
        'message': 'Invalid request format'
    }), 400

@app.errorhandler(500)
def internal_server_error(error):
    logger.error(f"Internal server error: {str(error)}")
    return jsonify({
        'error': 'Internal server error',
        'message': 'An unexpected error occurred'
    }), 500

if __name__ == '__main__':
    app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE
    
    try:
        # Initialize services
        logger.info("Initializing services...")
        initialize_services()
        
        logger.info("Starting Legal Document Risk Analyzer...")
        logger.info(f"Document AI: {'✓ Available' if _document_ai_client else '✗ Fallback mode'}")
        logger.info(f"Vertex AI: {'✓ Available' if _vertex_ai_available else '✗ Fallback mode'}")
        logger.info(f"Model: {MODEL_NAME}")
        logger.info(f"Risk patterns: {len(COMPREHENSIVE_RISK_PATTERNS)}")
        logger.info(f"Contract types: {len(CONTRACT_TYPE_PATTERNS) + 1}")
        logger.info(f"python-magic: {'✓' if MAGIC_AVAILABLE else '✗'}")
        
        # Run the application
        app.run(debug=True, host='0.0.0.0', port=5000, threaded=True)
        
    except Exception as e:
        logger.error(f"Failed to start application: {str(e)}")
        exit(1)